# src/core/fitz_extractor.py

import base64, os
import fitz  # PyMuPDF
from typing import List, Tuple, Optional
from loguru import logger
import tempfile
import pdfplumber
from pdf2docx import Converter
from docx import Document

from core.domain import (
    FitzSpan,
    FitzLine,
    FitzBlock,
    FitzPath,
    FitzPage,
    FitzDocument,
    FitzTableBlock
)

# from .formula_detector import FormulaDetector
from core.cid_normalizer import build_cid_maps, normalize_cids


# def page_has_table_lines(page: fitz.Page) -> bool:
#     """
#     Analyse instantanément (<1ms) les tracés vectoriels de la page.
#     Retourne True si la page possède au moins 3 lignes horizontales fines et larges,
#     ce qui est la signature physique obligatoire d'un tableau scientifique.
#     """
#     try:
#         drawings = page.get_drawings()
#         horizontal_lines = 0
#         for draw in drawings:
#             rect = draw.get("rect")
#             if not rect:
#                 continue
#             x0, y0, x1, y1 = rect
#             w = x1 - x0
#             h = y1 - y0
            
#             # Une ligne de tableau est large (ex: > 100 pt) et très fine (ex: < 3 pt)
#             if w > 100 and h < 3:
#                 horizontal_lines += 1
#                 if horizontal_lines >= 3:
#                     return True
#     except Exception:
#         pass
#     return False


class FitzExtractor:
    """
    Handles PDF data extraction using PyMuPDF (fitz).
    Converts PDF layout, text, and vector graphics into a structured FitzDocument.
    """

    def __init__(self, pdf_path: str, dpi: int = 150):
        """
        Args:
            pdf_path: Path to the target PDF file.
            dpi: Resolution for the page background PNG generation (150 is optimal for UI preview).
        """
        self.pdf_path = pdf_path
        self.dpi = dpi
        # self.cid_maps = None

    def extract_document(
        self,
        page_number: int = None,
        max_pages: int = None
    ) -> FitzDocument:
        """
        Parses the PDF and returns a structured FitzDocument.

        - If page_number is provided: extracts only that page.
        - If max_pages is provided: limits extraction to that number of pages.
        - If both are None: extracts the full document.
        """
        logger.info(f"Opening document: {self.pdf_path}")
        doc = fitz.open(self.pdf_path)
        fitz_doc = FitzDocument(path=self.pdf_path)

        total_pages = len(doc)

        if page_number is not None:
            # 1 page spécifique
            start = max(0, page_number - 1)
            end = min(start + 1, total_pages)

        else:
            # extraction batch (limit optionnelle)
            start = 0

            if max_pages is None:
                end = total_pages
            else:
                end = min(max_pages, total_pages)

        page_range = range(start, end)

        for page_num in page_range:
            page = doc[page_num]

            fitz_page = self._extract_page(
                page,
                page_num + 1,
                extract_tables=False
            )

            fitz_doc.pages.append(fitz_page)
            logger.info(f"Successfully extracted Page {page_num + 1}/{total_pages}")

        doc.close()
        return fitz_doc

    def _extract_page(self, page: fitz.Page, page_number: int, extract_tables: bool = False) -> FitzPage:
        """
        Extracts structural text blocks, vector paths, and the background image of a page.
        """

        # Initialisation paresseuse (lazy-loading) :
        # On lit le document parent de la page à la volée. Cela évite d'ouvrir le PDF deux fois
        # et empêche l'AttributeError si cette méthode est appelée seule (comme dans poc_render.py).
        if not hasattr(self, "cid_maps") or not self.cid_maps:
            # page.parent est l'objet fitz.Document d'origine du PDF !
            self.cid_maps = build_cid_maps(page.parent)

        page_w = page.rect.width
        page_h = page.rect.height

        # 1. Extract vector elements (Paths) first
        paths = self._extract_paths(page)

        # 2. Generate high-resolution background PNG
        png_b64 = self._generate_page_image_b64(page)

        # ← détecte les zones de tableaux et d'images
        table_blocks = []
        if extract_tables:
            table_blocks  = self._extract_tables(page)

        image_rects = self._extract_image_rects(page)
        skip_rects = [(tb.left, tb.top, tb.right, tb.bottom) for tb in table_blocks] + image_rects  # zones à ne pas couvrir


        # 3. Extract text blocks
        blocks = self._extract_text_blocks(page, page_number, paths, skip_rects )

        all_blocks = blocks + table_blocks

        # try:
        #     import pdfplumber
        #     with pdfplumber.open(self.pdf_path) as pdf:
        #         plumb_page = pdf.pages[page.number]
        #         plumb_words = plumb_page.extract_words(x_tolerance=1, y_tolerance=3)
        # except Exception:
        #     plumb_words = []

        # if plumb_words:
        #     detector = FormulaDetector(page_w, page_h, plumb_words)
        #     for block in blocks:
        #         if detector.is_formula(block):
        #             block.skip_translation = True

        return FitzPage(
            number=page_number,
            width=page_w,
            height=page_h,
            blocks=all_blocks,
            paths=paths,
            png_b64=png_b64
        )

    def _extract_paths(self, page: fitz.Page) -> List[FitzPath]:
        """
        Extracts drawings (fills, rectangles, strokes) to map physical visual decorations.
        """
        paths: List[FitzPath] = []
        try:
            drawings = page.get_drawings()
        except Exception as e:
            logger.warning(f"Failed to get drawings for page {page.number + 1}: {e}")
            return paths

        for draw in drawings:
            rect = draw.get("rect")
            if not rect:
                continue

            x0, y0, x1, y1 = rect
            w, h = x1 - x0, y1 - y0

            # Ignore extremely small lines or artifacts
            if w < 2.0 and h < 2.0:
                continue

            fill = draw.get("fill")
            color = draw.get("color")
            stroke_width = draw.get("width", 1.0) or 1.0

            # Convert normalized 0.0-1.0 RGB tuples to "rgb(r,g,b)"
            fill_css = f"rgb({int(fill[0]*255)},{int(fill[1]*255)},{int(fill[2]*255)})" if fill else None
            stroke_css = f"rgb({int(color[0]*255)},{int(color[1]*255)},{int(color[2]*255)})" if color else None

            paths.append(FitzPath(
                left=x0,
                top=y0,
                width=w,
                height=h,
                fill_color=fill_css,
                stroke_color=stroke_css,
                stroke_width=stroke_width
            ))

        return paths

    def _generate_page_image_b64(self, page: fitz.Page) -> str:
        """
        Renders the page to a PNG and encodes it in base64.
        """
        zoom = self.dpi / 72.0
        matrix = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        img_data = pix.tobytes("png")
        return base64.b64encode(img_data).decode("utf-8")

    def _extract_text_blocks(self, page: fitz.Page, page_number: int, paths: List[FitzPath], skip_rects=[]) -> List[FitzBlock]:
        """
        Extracts structural text layouts and maps background colors.
        """
        blocks: List[FitzBlock] = []
        
        # Extract rich dictionary layout from PyMuPDF
        text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        # text_dict = page.get_text("dict")
        
        block_id_counter = 0

        for b_dict in text_dict.get("blocks", []):
            # Type 0 is text. Ignore images/drawings blocks here
            if b_dict.get("type") != 0:
                continue

            x0, y0, x1, y1 = b_dict["bbox"]

            # Récupère le texte brut du bloc pour l'analyser
            raw_text = " ".join(s.get("text", "") for line in b_dict.get("lines", []) for s in line.get("spans", ""))
            if self._is_math_block(raw_text):
                continue  # On ignore la formule, elle reste propre sur le PNG d'origine !
            


            lines: List[FitzLine] = []

            for l_dict in b_dict.get("lines", []):
                lx0, ly0, lx1, ly1 = l_dict["bbox"]
                spans: List[FitzSpan] = []

                for s_dict in l_dict.get("spans", []):
                    sx0, sy0, sx1, sy1 = s_dict["bbox"]
                    raw_text = s_dict.get("text", "")
                     

                    # Extract styling
                    font = s_dict.get("font", "")
                    size = s_dict.get("size", 9.0)
                    flags = s_dict.get("flags", 0)
                    color_int = s_dict.get("color", 0)

                    # Extract RGB components from fitz integer color (sRGB)
                    r = (color_int >> 16) & 0xFF
                    g = (color_int >> 8) & 0xFF
                    b = color_int & 0xFF
                    color_css = f"rgb({r},{g},{b})"

                    # Detect Bold, Italic, Superscript
                    is_bold, is_italic = self._detect_font_style(font, flags)
                    is_sup = bool(flags & 1)  # Bit 0 indicates superscript
                    
                    text = normalize_cids(raw_text, font, self.cid_maps)

                    spans.append(FitzSpan(
                        text=text,
                        left=sx0,
                        top=sy0,
                        right=sx1,
                        bottom=sy1,
                        font_name=font,
                        font_size=size,
                        color=color_css,
                        is_bold=is_bold,
                        is_italic=is_italic,
                        is_sup=is_sup
                    ))

                if spans:
                    lines.append(FitzLine(
                        spans=spans,
                        left=lx0,
                        top=ly0,
                        right=lx1,
                        bottom=ly1
                    ))

            if not lines:
                continue

            # Compute actual layout line height
            line_height_ratio = self._compute_line_height(lines)

            # Determine if this block sits on top of a non-white vector background (e.g., Abstract grey boxes)
            bg_color = self._detect_background_color(x0, y0, x1, y1, paths)

            is_over_image = self._is_over_image(x0, y0, x1, y1, skip_rects)
            block = FitzBlock(
                block_id=block_id_counter,
                lines=lines,
                left=x0,
                top=y0,
                right=x1,
                bottom=y1,
                page_number=page_number,
                bg_color=bg_color,
                line_height_ratio=line_height_ratio
            )

            if is_over_image:
                block.bg_color = "transparent"  # ne cache pas l'image

            if self._is_in_skip_zone(x0, y0, x1, y1, skip_rects):
               continue  # ignore ce bloc — le PNG montre le tableau/image original
    

            blocks.append(block)
            block_id_counter += 1

        return blocks
    
    def _is_over_image(self, x0, y0, x1, y1, image_rects) -> bool:
        """Vérifie si ce bloc texte intersecte (chevauche) une zone d'image."""
        for ix0, iy0, ix1, iy1 in image_rects:
            # Calcul du chevauchement horizontal et vertical
            h_overlap = max(0, min(x1, ix1) - max(x0, ix0))
            v_overlap = max(0, min(y1, iy1) - max(y0, iy0))
            if h_overlap > 2 and v_overlap > 2:
                return True
        return False
    
    def _is_in_skip_zone(self, x0, y0, x1, y1, skip_rects) -> bool:
        """
        Retourne True si ce bloc de texte intersecte une zone de tableau ou d'image.
        Si intersecté, le texte ne sera pas dessiné pour laisser l'original propre.
        """
        for rx0, ry0, rx1, ry1 in skip_rects:
            # Vérification géométrique stricte d'intersection de deux bboxes
            h_overlap = max(0, min(x1, rx1) - max(x0, rx0))
            v_overlap = max(0, min(y1, ry1) - max(y0, ry0))
            # Si la surface de collision est réelle (> 2 points), il y a chevauchement
            if h_overlap > 2 and v_overlap > 2:
                return True
        return False


    def _extract_image_rects(self, page: fitz.Page) -> list:
        """
        Retourne les bboxes de toutes les images physiques de la page via l'API native.
        Ignore les images de fond décoratives (filigranes) de grande taille.
        """
        rects = []
        page_w = page.rect.width
        page_h = page.rect.height

        try:
            # get_image_info() renvoie une liste de dicts contenant la clé 'bbox' pour chaque image
            for img_info in page.get_image_info(hashes=True, xrefs=True):
                bbox = img_info.get("bbox")
                if not bbox:
                    continue
                
                # Conversion du tuple de coordonnées (x0, y0, x1, y1)
                x0, y0, x1, y1 = bbox
                w = x1 - x0
                h = y1 - y0
                
                # Optionnel : ignore les arrière-plans décoratifs plein écran (>95%)
                # if w > page_w * 0.95 or h > page_h * 0.95:
                #     continue
                
                rects.append((x0, y0, x1, y1))
        except Exception as e:
            logger.warning(f"Erreur lors de la détection des images : {e}")
       
        return rects
    

    def _extract_tables(self, page: fitz.Page) -> list:
        """
        Pipeline v5 :
        A. pdfplumber extract_words(x_tolerance=1) → mots avec bbox précises
        B. fitz get_text("dict") → styles par span (font_size, bold, italic, color)
        C. pdf2docx → détection zones tableau (is_real_table filter)
        D. Croiser cellules docx × mots pdfplumber → zone bbox globale
        E. Filtrer mots dans zone + enrichir avec styles fitz
        """
        

        table_blocks = []
        tmp_path = None

        try:
            # A. pdfplumber → mots avec bbox précises (gère la rotation)
            with pdfplumber.open(self.pdf_path) as pdf:
                p = pdf.pages[page.number]
                all_words = p.extract_words(x_tolerance=1, y_tolerance=3)

            # B. fitz → index des styles par position
            # Construit un dict {(x0_arrondi, top_arrondi): style}
            style_index = {}
            text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
            for b in text_dict.get("blocks", []):
                if b.get("type") != 0:
                    continue
                for line in b.get("lines", []):
                    for span in line.get("spans", []):
                        sx0, sy0, sx1, sy1 = span["bbox"]
                        font  = span.get("font", "")
                        size  = span.get("size", 8.5)
                        flags = span.get("flags", 0)
                        color_int = span.get("color", 0)
                        r = (color_int >> 16) & 0xFF
                        g = (color_int >> 8)  & 0xFF
                        b_val = color_int & 0xFF
                        is_bold, is_italic = self._detect_font_style(font, flags)
                        # Clé arrondie pour tolérance de matching
                        key = (round(sx0), round(sy0))
                        style_index[key] = {
                            "font_size": size,
                            "font_name": font,
                            "is_bold":   is_bold,
                            "is_italic": is_italic,
                            "color":     f"rgb({r},{g},{b_val})",
                        }

            # C. pdf2docx → détection tableaux
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp_path = tmp.name
            cv = Converter(self.pdf_path)
            cv.convert(tmp_path, pages=[page.number])
            cv.close()
            doc = Document(tmp_path)
            self._cleanup_temp_file(tmp_path)

            block_id = 10000
            seen_zones = []  # évite les doublons de zones

            for table in doc.tables:
                if not self._is_real_table(table):
                    continue

                # D. Textes cellules → matcher avec mots pdfplumber
                cell_texts = []
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip().split("\n")[0].strip()
                        if t and len(t) >= 3:
                            cell_texts.append(t.replace(" ", "").lower())

                matched = []
                for word in all_words:
                    wclean = word["text"].replace(" ", "").lower()
                    for ct in cell_texts:
                        if wclean and (wclean in ct or ct[:10] in wclean) and len(wclean) >= 3:
                            matched.append(word)
                            break

                if not matched:
                    continue

                zone = {
                    "x0":     min(w["x0"]     for w in matched) - 5,
                    "top":    min(w["top"]    for w in matched) - 5,
                    "x1":     max(w["x1"]     for w in matched) + 5,
                    "bottom": max(w["bottom"] for w in matched) + 5,
                }

                # Dédoublonnage : ignore si zone trop similaire à une déjà vue
                is_duplicate = any(
                    abs(zone["x0"] - z["x0"]) < 10 and
                    abs(zone["top"] - z["top"]) < 10 and
                    abs(zone["x1"] - z["x1"]) < 10 and
                    abs(zone["bottom"] - z["bottom"]) < 10
                    for z in seen_zones
                )
                if is_duplicate:
                    continue
                seen_zones.append(zone)

                # E. Filtrer mots dans la zone + enrichir avec styles fitz
                table_words = []
                for w in all_words:
                    if not (w["x0"] >= zone["x0"] and w["x1"] <= zone["x1"] and
                            w["top"] >= zone["top"] and w["bottom"] <= zone["bottom"]):
                        continue

                    # Cherche le style fitz le plus proche par position
                    key = (round(w["x0"]), round(w["top"]))
                    style = style_index.get(key)
                    if not style:
                        # Tolérance ±3px
                        for dx in range(-3, 4):
                            for dy in range(-3, 4):
                                style = style_index.get((key[0]+dx, key[1]+dy))
                                if style:
                                    break
                            if style:
                                break

                    font_for_resolve = style["font_name"] if style else "Unknown"
                    cleaned_text = normalize_cids(w["text"], font_for_resolve, self.cid_maps)
                    table_words.append({
                        "text":      cleaned_text,
                        "x0":        w["x0"],
                        "top":       w["top"],
                        "x1":        w["x1"],
                        "bottom":    w["bottom"],
                        "font_size": style["font_size"] if style else 8.5,
                        "is_bold":   style["is_bold"]   if style else False,
                        "is_italic": style["is_italic"] if style else False,
                        "color":     style["color"]     if style else "rgb(0,0,0)",
                    })

                if not table_words:
                    continue

                table_blocks.append(FitzTableBlock(
                    block_id=block_id,
                    left=zone["x0"],
                    top=zone["top"],
                    right=zone["x1"],
                    bottom=zone["bottom"],
                    page_number=page.number + 1,
                    words=table_words,
                ))
                block_id += 1
                logger.info(f"  Table {block_id-10000}: {len(table_words)} mots, zone top={zone['top']:.0f}")

            logger.info(f"Page {page.number+1} → {len(table_blocks)} tableau(x)")

        except Exception as e:
            self._cleanup_temp_file(tmp_path)
            import traceback; traceback.print_exc()
            logger.warning(f"Table extraction failed p{page.number+1}: {e}")

        return table_blocks


    def _is_real_table(self, table) -> bool:
        """Filtre les faux positifs : texte en colonnes détecté comme tableau."""
        if len(table.rows) < 2 or len(table.columns) < 2:
            return False
        first_row_texts = [c.text.strip() for c in table.rows[0].cells]
        long_cells = sum(1 for t in first_row_texts if len(t) > 80)
        if long_cells >= len(table.columns) - 1:
            return False
        return True

    def _cleanup_temp_file(self, path: Optional[str]):
        """Supprime proprement les fichiers temporaires pour éviter les fuites de ressources."""
        if path and os.path.exists(path):
            try:
                os.remove(path)
            except Exception:
                pass

            
       
    
    def _is_math_block(self, text: str) -> bool:
        """
        Détecte si un bloc de texte représente une formule mathématique ou une équation.
        """
        t = text.strip()
        if not t:
            return False

        # Heuristique 1 : Se termine par un numéro d'équation entre parenthèses, ex: "(2)" ou "(4)"
        import re
        if re.search(r'\(\d+\)$', t) or re.search(r'\(Eq\.\s*\d+\)$', t):
            return True

        # Heuristique 2 : Contient des symboles mathématiques complexes
        math_symbols = {'∑', '∫', '∏', '√', '±', '−', '×', '÷', '≠', '≤', '≥', '≈', '≡', 'λ', 'μ', 'π', 'σ', '∂', '∆', '∇'}
        if any(sym in t for sym in math_symbols):
            return True

        # Heuristique 3 : Formule courte contenant une égalité (ex: "aij = 1/aji")
        if " = " in t and len(t) < 80:
            return True

        return False
    
    def _detect_font_style(self, font_name: str, flags: int) -> Tuple[bool, bool]:
        """
        Uses both font name flags and binary bitwise flags to robustly spot bold/italic variations.
        """
        font_lower = font_name.lower()
        
        # Check standard PyMuPDF flags
        is_bold = bool(flags & 16) or any(x in font_lower for x in ["bold", "black", "heavy", "-b"])
        is_italic = bool(flags & 2) or any(x in font_lower for x in ["italic", "oblique", "-i"])

        return is_bold, is_italic

    def _compute_line_height(self, lines: List[FitzLine]) -> float:
        """
        Computes the visual line spacing ratio (useful for pixel-perfect CSS rendering).
        """
        if len(lines) < 2:
            return 1.1

        spacings = []
        for i in range(1, len(lines)):
            delta = lines[i].top - lines[i-1].top
            if delta > 0:
                spacings.append(delta)

        if not spacings:
            return 1.1

        avg_spacing = sum(spacings) / len(spacings)
        dominant_fs = max(s.font_size for line in lines for s in line.spans) if lines else 9.0
        
        ratio = avg_spacing / max(dominant_fs, 1.0)
        return min(max(ratio, 0.9), 2.0)  # Clamp values to avoid weird layouts

    def _detect_background_color(self, x0: float, y0: float, x1: float, y1: float, paths: List[FitzPath]) -> str:
        """
        Spatially checks if any solid drawing sits directly under this block's center coordinate.
        """
        cx = (x0 + x1) / 2.0
        cy = (y0 + y1) / 2.0

        for path in paths:
            if not path.fill_color:
                continue

            # Check bounding box containment
            px0 = path.left
            py0 = path.top
            px1 = path.left + path.width
            py1 = path.top + path.height

            if (px0 <= cx <= px1) and (py0 <= cy <= py1):
                # Récupère les valeurs RGB de la chaîne "rgb(r,g,b)"
                try:
                    parts = [int(v) for v in path.fill_color.replace("rgb(", "").replace(")", "").split(",")]
                    # Si toutes les composantes sont > 240 (gris très clair / blanc), on ignore
                    if all(c > 240 for c in parts):
                        continue
                except ValueError:
                    continue
                
                return path.fill_color

        return "white"